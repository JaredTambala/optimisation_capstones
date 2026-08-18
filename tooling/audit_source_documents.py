"""Audit the CAP-001 decision configuration against the approved DOCX sources.

This is an independent control-side audit. It intentionally reads the frozen
source documents rather than trusting generated schemas or documentation.
Install the optional audit dependencies with `pip install -e '.[audit]'`.
"""

from __future__ import annotations

import re
import sys

from tooling.contract_runtime import (
    EXPECTED_RAW_FILES,
    ROOT,
    ContractError,
    load_config,
    sha256_path,
)


SPEC_PATH = ROOT / "CAP-001_Tier-N_End-to-End_Cost_Model_Design_and_Dataset_Generation_Specification_v0.3.docx"
STANDARD_PATH = ROOT / "Optimisation_Search_and_Decision_Intelligence_Capstone_Control_Standard_v0.2.docx"
CN003_PATH = ROOT / "CAP-001_CONTRACT_PRICING_LABEL_CHANGE_NOTES.md"

# CN-003 removes a non-domain labelling column from the effective contract.
# The frozen v0.3 DOCX remains unchanged and is amended by the versioned note.
CONTROLLED_SOURCE_FIELD_OMISSIONS = {
    "supply_contracts.csv": {"pricing_method"},
}


def _cell(cell) -> str:
    return " ".join(cell.text.split())


def audit() -> dict[str, int]:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - environment guidance
        raise ContractError("python-docx is required; install the 'audit' extra") from exc

    config = load_config()
    if not CN003_PATH.is_file():
        raise ContractError("CN-003 is required for the controlled supply-contract field omission")
    expected_spec_hash = config["document_control"]["capstone_specification"]["sha256"]
    expected_standard_hash = config["document_control"]["control_standard"]["sha256"]
    if sha256_path(SPEC_PATH) != expected_spec_hash or sha256_path(STANDARD_PATH) != expected_standard_hash:
        raise ContractError("governing source document hash mismatch")

    specification = Document(SPEC_PATH)
    standard = Document(STANDARD_PATH)

    # Tables 25-50 are the 26 raw contracts, in the release order frozen by v0.3.
    type_map = {"string": "string", "integer": "integer", "decimal": "number", "date": "string", "timestamp": "string", "boolean": "boolean", "category": "string"}
    raw_fields = 0
    for file_name, table_index in zip(EXPECTED_RAW_FILES, range(25, 51), strict=True):
        table = specification.tables[table_index]
        source_fields = [_cell(row.cells[0]) for row in table.rows[1:]]
        source_types = [_cell(row.cells[1]) for row in table.rows[1:]]
        configured = config["raw_contracts"][file_name]["columns"]
        configured_fields = [field["name"] for field in configured]
        configured_types = [field["type"] for field in configured]
        effective_source_fields = [
            field
            for field in source_fields
            if field not in CONTROLLED_SOURCE_FIELD_OMISSIONS.get(file_name, set())
        ]
        effective_source_types = [
            field_type
            for field, field_type in zip(source_fields, source_types, strict=True)
            if field not in CONTROLLED_SOURCE_FIELD_OMISSIONS.get(file_name, set())
        ]
        if effective_source_fields != configured_fields:
            raise ContractError(f"{file_name}: field names/order differ from source table {table_index}")
        expected_types = [type_map[value] for value in effective_source_types]
        if expected_types != configured_types:
            raise ContractError(f"{file_name}: field types differ from source table {table_index}")
        for source_type, field in zip(effective_source_types, configured, strict=True):
            if source_type == "date" and field.get("format") != "date":
                raise ContractError(f"{file_name}.{field['name']}: date format missing")
            if source_type == "timestamp" and field.get("format") != "date-time":
                raise ContractError(f"{file_name}.{field['name']}: timestamp format missing")
        raw_fields += len(configured_fields)

    # CAP-specific output names are table 58; the common evaluation minimum is table 7.
    spec_outputs = [_cell(row.cells[0]) for row in specification.tables[58].rows[1:]]
    missing_spec_outputs = set(spec_outputs) - set(config["output_contracts"])
    if missing_spec_outputs:
        raise ContractError(f"CAP output contracts missing: {sorted(missing_spec_outputs)}")
    common_text = " ".join(_cell(cell) for row in standard.tables[7].rows for cell in row.cells)
    common_outputs = set(re.findall(r"artifacts/evaluation/([A-Za-z0-9_.-]+)", common_text))
    missing_common_outputs = common_outputs - set(config["output_contracts"])
    if missing_common_outputs:
        raise ContractError(f"common output contracts missing: {sorted(missing_common_outputs)}")

    # Rubric, solver statuses and runtime budgets are frozen machine controls.
    source_rubric = [(_cell(row.cells[0]), int(_cell(row.cells[1]))) for row in specification.tables[62].rows[1:]]
    configured_rubric = [(item["category"], item["points"]) for item in config["assessment"]["rubric"]]
    if source_rubric != configured_rubric:
        raise ContractError("CAP-001 rubric differs from source table 62")
    source_statuses = [_cell(row.cells[0]) for row in specification.tables[20].rows[1:]]
    if source_statuses != config["solution_statuses"]:
        raise ContractError("solution status vocabulary differs from source table 20")
    budgets = config["runtime_budgets"]
    if not (
        budgets["miniature_fixture_seconds"] == 120
        and budgets["baseline_per_scenario_seconds"] == 300
        and budgets["recursive_base"]["maximum_starts"] == 3
        and budgets["recursive_base"]["seconds_per_start"] == 1200
        and budgets["recursive_scenario_reoptimisation_seconds"] == 900
    ):
        raise ContractError("runtime budgets differ from source table 61")

    return {
        "governing_documents": 2,
        "raw_contracts": len(EXPECTED_RAW_FILES),
        "raw_fields": raw_fields,
        "cap_specific_outputs": len(spec_outputs),
        "common_outputs": len(common_outputs),
        "rubric_categories": len(source_rubric),
        "solution_statuses": len(source_statuses),
    }


def main() -> int:
    try:
        summary = audit()
    except (ContractError, KeyError, IndexError, ValueError) as exc:
        print(f"Source-document audit failed: {exc}", file=sys.stderr)
        return 1
    print("Source-document audit passed:")
    for name, value in summary.items():
        print(f"  {name}: {value}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
